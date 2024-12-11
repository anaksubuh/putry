import streamlit as st
from docx import Document
import nltk
from nltk.corpus import wordnet
from nltk.tokenize import word_tokenize, sent_tokenize
import random
import io
from googletrans import Translator

# Pastikan resource NLTK telah diunduh
nltk.download('punkt')
nltk.download('wordnet')

def get_synonyms(word):
    synonyms = set()
    for syn in wordnet.synsets(word):
        for lemma in syn.lemmas():
            synonyms.add(lemma.name())
    if word in synonyms:
        synonyms.remove(word)
    return list(synonyms)

def paraphrase_sentence(sentence):
    words = word_tokenize(sentence)
    new_sentence = []
    for word in words:
        synonyms = get_synonyms(word)
        if synonyms:
            new_sentence.append(random.choice(synonyms))
        else:
            new_sentence.append(word)
    return ' '.join(new_sentence)

def paraphrase_text(text):
    sentences = sent_tokenize(text)
    new_text = [paraphrase_sentence(sentence) for sentence in sentences]
    return ' '.join(new_text)

def create_word_document(text):
    document = Document()
    document.add_paragraph(text)
    
    # Save to an in-memory file
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    return buffer

def read_text_from_word(file):
    document = Document(file)
    text = '\n'.join([para.text for para in document.paragraphs])
    return text

# Streamlit App
st.set_page_config(
    page_title='PUTRY CANTIK',
    page_icon='logo.png',
    initial_sidebar_state='expanded',
    menu_items={
        'Get Help': 'https://www.extremelycoolapp.com/help',
        'Report a bug': 'https://www.extremelycoolapp.com/bug',
        'About': '# This is a header. This is an *extremely* cool app!'
    }
)

st.title('<3 Putry Paraphrasing Tool <3')

option = st.selectbox('Choose input method:', ('Copy-Paste Text', 'Upload File'))

def translate_text(text, source_lang="auto", target_lang="en"):
    """Function to translate text using Google Translate"""
    translator = Translator()
    translation = translator.translate(text, src=source_lang, dest=target_lang)
    return translation.text

def read_text_from_word(uploaded_file):
    """Function to read text from a Word document"""
    from docx import Document
    doc = Document(uploaded_file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

if option == 'Translator':
    option = st.selectbox('Choose input method:', ('Copy-Paste Text', 'Upload File'))

    if option == 'Copy-Paste Text':
        input_text = st.text_area('Enter text to translate:', height=300)
    else:
        uploaded_file = st.file_uploader('Upload a file:', type=['txt', 'docx'])
        if uploaded_file:
            if uploaded_file.type == 'text/plain':
                input_text = uploaded_file.read().decode('utf-8')
            elif uploaded_file.type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                input_text = read_text_from_word(uploaded_file)
            st.text_area('File content:', input_text, height=300, disabled=True)

    target_lang = st.selectbox('Choose target language:', ('en', 'id', 'es', 'fr', 'de', 'it'))

    if st.button('Translate Text'):
        if input_text:
            translated_text = translate_text(input_text, source_lang="auto", target_lang=target_lang)
            st.text_area('Translated text:', translated_text, height=300)

            # For downloading the translated text as a Word document
            from io import BytesIO
            from docx import Document

            def create_word_document(translated_text):
                doc = Document()
                doc.add_paragraph(translated_text)
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                return buffer

            buffer = create_word_document(translated_text)
            st.download_button(
                label='Download Translated Document',
                data=buffer,
                file_name='translated_document.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            st.success('Translated text is ready to download.')
        else:
            st.error('Please enter or upload text to translate.')

    st.warning('This translation is not 100% accurate, please verify manually.', icon="⚠️")
    st.info('Translation tool by Casper', icon="ℹ️")

elif option == 'Parafrase':
    option = st.selectbox('Choose input method:', ('Copy-Paste Text', 'Upload File'))

    if option == 'Copy-Paste Text':
        input_text = st.text_area('Enter text to paraphrase:', height=300)
    else:
        uploaded_file = st.file_uploader('Upload a file:', type=['txt', 'docx'])
        if uploaded_file:
            if uploaded_file.type == 'text/plain':
                input_text = uploaded_file.read().decode('utf-8')
            elif uploaded_file.type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                input_text = read_text_from_word(uploaded_file)
            st.text_area('File content:', input_text, height=300, disabled=True)

    if st.button('Paraphrase Text'):
        if input_text:
            paraphrased_text = paraphrase_text(input_text)
            st.text_area('Paraphrased text:', paraphrased_text, height=300)
            
            buffer = create_word_document(paraphrased_text)
            st.download_button(
                label='Download Paraphrased Document',
                data=buffer,
                file_name='sukses_analyzer.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            st.success('Paraphrased text is ready to download.')
        else:
            st.error('Please enter or upload text to paraphrase.')

    st.warning('ROBOT INI TIDAK 100% AKURAT , MOHON DI CHEK MANUAL UNTUK HASILNYA', icon="⚠️")
    st.info('Robot by casper', icon="ℹ️")
